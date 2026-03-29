"""
TenderWriter — OnlyOffice Integration API

Provides document serving, callback handling, and RAG indexing
for OnlyOffice Document Server integration.
"""

from __future__ import annotations

import hashlib
import hmac
import io
import secrets
import time
import jwt
from datetime import datetime, timedelta, timezone
from urllib.parse import urlencode

import httpx
import structlog
from docx import Document as DocxDocument
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import JSONResponse, StreamingResponse
from minio import Minio
from minio.error import S3Error
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.db.database import get_db
from app.db.redis import redis_client
from app.models import Proposal, ProposalSection, ContentBlock
from app.api.auth import get_current_user, UserResponse
from app.utils.naming import sanitize_name, get_structured_minio_path
from app.services.compliance_observability import sync_requirement_compliance_and_gate
from app.services.kpi_reason_engine import (
    build_proposal_section_updated_event_payload,
    publish_domain_event,
    sync_tender_and_publish_event,
)

logger = structlog.get_logger()

router = APIRouter()


class MinioDocumentStore:
    """MinIO-based document storage for OnlyOffice documents.

    Replaces in-memory dict storage with persistent MinIO storage.
    Includes TTL-based cleanup for old documents.

    All synchronous minio-client calls are wrapped with asyncio.to_thread()
    so they never block the event loop.
    """

    def __init__(self):
        self.bucket_name = settings.minio_bucket
        self.client = Minio(
            settings.minio_endpoint,
            access_key=settings.minio_access_key,
            secret_key=settings.minio_secret_key,
            secure=settings.minio_secure,
        )

    def _ensure_bucket(self):
        """Create bucket if it doesn't exist."""
        try:
            if not self.client.bucket_exists(self.bucket_name):
                self.client.make_bucket(self.bucket_name)
                logger.info("Created MinIO bucket", bucket=self.bucket_name)
        except S3Error as e:
            logger.error("Failed to ensure bucket", error=str(e))

    def _save_sync(self, doc_key: str, docx_bytes: bytes) -> None:
        data = io.BytesIO(docx_bytes)
        self.client.put_object(
            self.bucket_name,
            doc_key,
            data,
            length=len(docx_bytes),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    async def save(self, doc_key: str, docx_bytes: bytes) -> None:
        """Save document to MinIO (non-blocking)."""
        import asyncio
        try:
            await asyncio.to_thread(self._save_sync, doc_key, docx_bytes)
            logger.info("Document saved to MinIO", doc_key=doc_key, size=len(docx_bytes))
        except S3Error as e:
            logger.error("Failed to save document to MinIO", doc_key=doc_key, error=str(e))
            raise

    def _get_sync(self, doc_key: str) -> bytes | None:
        response = self.client.get_object(self.bucket_name, doc_key)
        data = response.read()
        response.close()
        response.release_conn()
        return data

    async def get(self, doc_key: str) -> bytes | None:
        """Retrieve document from MinIO (non-blocking)."""
        import asyncio
        try:
            return await asyncio.to_thread(self._get_sync, doc_key)
        except S3Error as e:
            if e.code == "NoSuchKey":
                return None
            logger.error("Failed to get document from MinIO", doc_key=doc_key, error=str(e))
            raise

    async def delete(self, doc_key: str) -> bool:
        """Delete document from MinIO (non-blocking)."""
        import asyncio
        try:
            await asyncio.to_thread(self.client.remove_object, self.bucket_name, doc_key)
            logger.info("Document deleted from MinIO", doc_key=doc_key)
            return True
        except S3Error as e:
            logger.error("Failed to delete document from MinIO", doc_key=doc_key, error=str(e))
            return False

    async def cleanup_old_documents(self, max_age_hours: int = 24) -> int:
        """Remove documents older than max_age_hours (non-blocking).

        Returns count of deleted documents.
        """
        import asyncio
        from datetime import timedelta

        def _cleanup_sync():
            objects = self.client.list_objects(self.bucket_name)
            deleted_count = 0
            cutoff_time = datetime.now(timezone.utc) - timedelta(hours=max_age_hours)

            for obj in objects:
                last_modified = obj.last_modified
                if last_modified and last_modified.astimezone(timezone.utc) < cutoff_time:
                    self.client.remove_object(self.bucket_name, obj.object_name)
                    deleted_count += 1
                    logger.info("Old document deleted", doc_key=obj.object_name, last_modified=last_modified)

            return deleted_count

        try:
            return await asyncio.to_thread(_cleanup_sync)
        except S3Error as e:
            logger.error("Failed to cleanup old documents", error=str(e))
            return 0


# Singleton instance
_document_store = MinioDocumentStore()


# ── Schemas ──


class OnlyOfficeConfigResponse(BaseModel):
    """Full configuration for OnlyOffice DocEditor initialization."""
    config: dict
    token: str
    onlyoffice_url: str


# ── Helpers ──


def _build_config_dict(
    doc_key: str, 
    title: str, 
    file_url: str, 
    callback_url: str,
    user_id: str = "guest",
    user_name: str = "TenderWriter User"
) -> dict:
    """Standardize OnlyOffice configuration object structure."""
    return {
        "document": {
            "fileType": "docx",
            "key": doc_key,
            "title": title,
            "url": file_url,
            "permissions": {
                "comment": True,
                "copy": True,
                "download": True,
                "edit": True,
                "print": True,
                "fillForms": True,
                "chat": True,
            }
        },
        "documentType": "word",
        "width": "100%",
        "height": "100%",
        "editorConfig": {
            "callbackUrl": callback_url,
            "lang": "it",
            "mode": "edit",
            "user": {
                "id": user_id,
                "name": user_name,
            },
            "customization": {
                "autosave": True,
                "comments": True,
                "compactHeader": True,
                "compactToolbar": False,
                "feedback": False,
                "forcesave": True,
                "help": True,
                "hideRightMenu": True,
                "hideRulers": False,
                # Force a deterministic UI theme across proposal/library/create editors.
                "uiTheme": settings.onlyoffice_ui_theme,
                "logo": {
                    "image": "",
                    "visible": False,
                },
                "toolbarNoTabs": False,
            },
        },
    }


def _get_minio_object_name(doc_type: str, entity_id: int | str) -> str:
    """Legacy helper. Structured paths preferred."""
    return f"{doc_type}/{entity_id}.docx"


def _generate_document_key(proposal_id: int, section_id: int, updated_at: datetime | None) -> str:
    """Generate a dynamic document key for OnlyOffice based on version.
    
    OnlyOffice requires a new key when the document is changed.
    We use the updated_at timestamp as a version identifier.
    """
    ts = int(updated_at.timestamp()) if updated_at else 0
    raw = f"p{proposal_id}_s{section_id}_{ts}"
    return hmac.new(
        settings.onlyoffice_jwt_secret.encode("utf-8"),
        raw.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()[:32]


def _generate_library_document_key(block_id: int, updated_at: datetime | None) -> str:
    """Generate a dynamic document key for a library block."""
    ts = int(updated_at.timestamp()) if updated_at else 0
    raw = f"lib_{block_id}_{ts}"
    return hmac.new(
        settings.onlyoffice_jwt_secret.encode("utf-8"),
        raw.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()[:32]


def _generate_create_document_key() -> str:
    """Generate a unique document key for a new document creation.
    We keep this one unique as it's for a temporary 'new' doc.
    """
    raw = f"create_{time.time_ns()}_{secrets.token_hex(8)}"
    return hmac.new(
        settings.onlyoffice_jwt_secret.encode("utf-8"),
        raw.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()[:32]


def _build_file_signature(doc_key: str, expires_at: int) -> str:
    payload = f"{doc_key}:{expires_at}"
    return hmac.new(
        settings.onlyoffice_jwt_secret.encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()


def _build_download_token(*, doc_key: str, user_id: int | str) -> str:
    return jwt.encode(
        {
            "sub": str(user_id),
            "doc_key": doc_key,
            "exp": datetime.now(timezone.utc) + timedelta(seconds=settings.onlyoffice_file_token_ttl_seconds),
        },
        settings.onlyoffice_jwt_secret,
        algorithm="HS256",
    )


def _verify_download_token(doc_key: str, metadata: dict, download_token: str) -> bool:
    try:
        payload = jwt.decode(download_token, settings.onlyoffice_jwt_secret, algorithms=["HS256"])
    except jwt.InvalidTokenError:
        return False
    if str(payload.get("doc_key") or "") != doc_key:
        return False
    owner_user_id = metadata.get("owner_user_id")
    return not owner_user_id or str(owner_user_id) == str(payload.get("sub") or "")


def _validate_callback_token(payload_token: str | None, doc_key: str) -> bool:
    if not payload_token:
        return False
    try:
        decoded = jwt.decode(payload_token, settings.onlyoffice_jwt_secret, algorithms=["HS256"])
    except jwt.InvalidTokenError:
        return False
    direct_key = str(decoded.get("key") or "")
    document_key = str((decoded.get("document") or {}).get("key") or "")
    return doc_key in {direct_key, document_key}


def _build_signed_file_url(doc_key: str, *, download_token: str) -> str:
    expires_at = int(time.time()) + settings.onlyoffice_file_token_ttl_seconds
    query = urlencode(
        {
            "expires": expires_at,
            "signature": _build_file_signature(doc_key, expires_at),
            "download_token": download_token,
        }
    )
    return f"{settings.backend_public_url}/api/onlyoffice/files/{doc_key}?{query}"


def _verify_file_signature(doc_key: str, expires_at: int, signature: str) -> bool:
    if expires_at < int(time.time()):
        return False
    expected = _build_file_signature(doc_key, expires_at)
    return hmac.compare_digest(expected, signature)


def _section_content_to_text(content: dict | None) -> str:
    """Extract plain text from section content (TipTap JSON format)."""
    if not content or not isinstance(content, dict):
        return ""
    
    if content.get("type") == "doc" and isinstance(content.get("content"), list):
        paragraphs = []
        for node in content["content"]:
            if isinstance(node, dict) and isinstance(node.get("content"), list):
                text_parts = [c.get("text", "") for c in node["content"] if isinstance(c, dict)]
                paragraphs.append("".join(text_parts))
        return "\n\n".join(paragraphs)
    
    if isinstance(content.get("text"), str):
        return content["text"]
    
    return ""


def _create_docx_from_text(text: str, title: str = "") -> bytes:
    """Create a .docx file from plain text."""
    doc = DocxDocument()
    
    if title:
        doc.add_heading(title, level=1)
    
    if text.strip():
        for paragraph in text.split("\n\n"):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                for run in p.runs:
                    run.font.size = Pt(11)
    else:
        doc.add_paragraph("")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract plain text from a .docx file."""
    doc = DocxDocument(io.BytesIO(docx_bytes))
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)
    return "\n\n".join(paragraphs)


def _text_to_tiptap_content(text: str) -> dict:
    """Convert plain text to TipTap-compatible JSON."""
    return {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": para}],
            }
            for para in text.split("\n\n")
            if para.strip()
        ] or [{"type": "paragraph", "content": [{"type": "text", "text": ""}]}],
    }


# Redis-backed metadata persistence
# We store keys in Redis to survive worker restarts
import json

async def save_session_metadata(key: str, metadata: dict, ttl: int = 86400):
    """Store session metadata in Redis with a TTL (default 24h)."""
    await redis_client.setex(f"oo_session:meta:{key}", ttl, json.dumps(metadata))
    
    # Also track the "last active key" for specific entities to support force-save lookup
    if "section_id" in metadata:
        entity_id = f"section:{metadata['proposal_id']}:{metadata['section_id']}"
        await redis_client.setex(f"oo_session:active:{entity_id}", ttl, key)
    elif "block_id" in metadata:
        entity_id = f"block:{metadata['block_id']}"
        await redis_client.setex(f"oo_session:active:{entity_id}", ttl, key)

async def get_session_metadata(key: str) -> dict | None:
    """Retrieve session metadata from Redis."""
    data = await redis_client.get(f"oo_session:meta:{key}")
    if not data:
        return None
    try:
        return json.loads(data)
    except (json.JSONDecodeError, ValueError):
        logger.warning("Corrupted session metadata in Redis", key=key)
        return None

async def get_active_key(entity_type: str, entity_id: str) -> str | None:
    """Retrieve the active session key for an entity."""
    return await redis_client.get(f"oo_session:active:{entity_type}:{entity_id}")


# ── Routes ──


@router.get("/document/proposal/{proposal_id}/{section_id}", response_model=OnlyOfficeConfigResponse)
async def get_document_config(
    proposal_id: int,
    section_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Generate a document config for OnlyOffice editor.
    Creates a .docx from the section content and returns config to initialize the editor.
    """
    # Fetch the section with proposal and tender
    result = await db.execute(
        select(ProposalSection)
        .where(
            ProposalSection.id == section_id,
            ProposalSection.proposal_id == proposal_id,
        )
        .options(
            selectinload(ProposalSection.proposal).selectinload(Proposal.tender)
        )
    )
    section = result.scalar_one_or_none()
    if not section:
        raise HTTPException(status_code=404, detail="Section not found")
    
    proposal = section.proposal
    tender = proposal.tender

    # Extract text from section content
    text = _section_content_to_text(section.content)
    
    # Create .docx
    docx_bytes = _create_docx_from_text(text, title=section.title)
    
    # Determine user prefix for folder structure
    user_prefix = current_user.email.split('@')[0] if current_user.email else "unknown"

    # Generate unique key including version info
    doc_key = _generate_document_key(proposal_id, section_id, section.updated_at or section.created_at)
    
    # Store in MinIO using a STABLE STRUCTURED path
    object_name = get_structured_minio_path(
        user_prefix=user_prefix,
        tender_title=tender.title,
        tender_id=tender.id,
        section_title=section.title,
        section_id=section.id
    )
    await _document_store.save(object_name, docx_bytes)
    
    # Store metadata in Redis
    await save_session_metadata(doc_key, {
        "type": "section",
        "proposal_id": proposal_id,
        "section_id": section_id,
        "object_name": object_name,
        "owner_user_id": current_user.id,
    })
    
    # Build URLs
    download_token = _build_download_token(doc_key=doc_key, user_id=current_user.id)
    file_url = _build_signed_file_url(doc_key, download_token=download_token)
    callback_url = f"{settings.backend_public_url}/api/onlyoffice/callback"
    
    logger.info(
        "OnlyOffice document config generated",
        proposal_id=proposal_id,
        section_id=section_id,
        doc_key=doc_key,
        file_url=file_url,
    )
    
    config = _build_config_dict(
        doc_key=doc_key,
        title=f"{section.title}.docx",
        file_url=file_url,
        callback_url=callback_url,
        user_id=str(current_user.id),
        user_name=current_user.name
    )
    
    token = jwt.encode(config, settings.onlyoffice_jwt_secret, algorithm="HS256")
    
    return OnlyOfficeConfigResponse(
        config=config,
        token=token,
        onlyoffice_url=settings.onlyoffice_url,
    )


@router.get("/document/create", response_model=OnlyOfficeConfigResponse)
async def get_create_document_config(
    current_user: UserResponse = Depends(get_current_user),
):
    """
    Generate a document config for a NEW document (creation flow).
    """
    # Create empty .docx
    docx_bytes = _create_docx_from_text("", title="Nuovo Blocco")
    
    # Generate unique key
    doc_key = _generate_create_document_key()
    
    # Store in MinIO
    await _document_store.save(doc_key, docx_bytes)
    
    # Build URLs
    download_token = _build_download_token(doc_key=doc_key, user_id=current_user.id)
    file_url = _build_signed_file_url(doc_key, download_token=download_token)
    callback_url = f"{settings.backend_public_url}/api/onlyoffice/callback"
    
    # Store metadata in Redis
    await save_session_metadata(doc_key, {
        "type": "create",
        "object_name": doc_key,  # For 'create', the doc_key itself is unique enough
        "owner_user_id": current_user.id,
    })
    
    config = _build_config_dict(
        doc_key=doc_key,
        title="Nuovo Documento.docx",
        file_url=file_url,
        callback_url=callback_url,
        user_id=str(current_user.id),
        user_name=current_user.name
    )
    
    token = jwt.encode(config, settings.onlyoffice_jwt_secret, algorithm="HS256")
    
    return OnlyOfficeConfigResponse(
        config=config,
        token=token,
        onlyoffice_url=settings.onlyoffice_url,
    )


@router.get("/document/library/{block_id}", response_model=OnlyOfficeConfigResponse)
async def get_library_document_config(
    block_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Generate a document config for OnlyOffice editor for a Content Library block.
    """
    # Fetch the block
    result = await db.execute(
        select(ContentBlock).where(ContentBlock.id == block_id)
    )
    block = result.scalar_one_or_none()
    if not block:
        raise HTTPException(status_code=404, detail="Content block not found")

    # ContentBlock uses raw strings, but we can treat it similarly
    text = block.content or ""
    
    # Create .docx
    docx_bytes = _create_docx_from_text(text, title=block.title)
    
    # Generate unique key including version info
    doc_key = _generate_library_document_key(block_id, block.updated_at or block.created_at)
    
    # Store in MinIO using a STABLE path
    object_name = _get_minio_object_name("block", block_id)
    await _document_store.save(object_name, docx_bytes)
    
    # Store metadata in Redis
    await save_session_metadata(doc_key, {
        "type": "library_block",
        "block_id": block_id,
        "object_name": object_name,
        "owner_user_id": current_user.id,
    })
    
    # Build URLs
    download_token = _build_download_token(doc_key=doc_key, user_id=current_user.id)
    file_url = _build_signed_file_url(doc_key, download_token=download_token)
    callback_url = f"{settings.backend_public_url}/api/onlyoffice/callback"
    
    logger.info(
        "OnlyOffice document config generated for library block",
        block_id=block_id,
        doc_key=doc_key,
        file_url=file_url,
    )
    
    config = _build_config_dict(
        doc_key=doc_key,
        title=f"{block.title}.docx",
        file_url=file_url,
        callback_url=callback_url,
        user_id=str(current_user.id),
        user_name=current_user.name
    )
    
    token = jwt.encode(config, settings.onlyoffice_jwt_secret, algorithm="HS256")
    
    return OnlyOfficeConfigResponse(
        config=config,
        token=token,
        onlyoffice_url=settings.onlyoffice_url,
    )


@router.get("/files/{doc_key}")
async def serve_document(doc_key: str, expires: int, signature: str, download_token: str):
    """Serve a .docx file to OnlyOffice Document Server."""
    if not _verify_file_signature(doc_key, expires, signature):
        raise HTTPException(status_code=403, detail="Invalid or expired file token")

    # Find metadata to get the actual MinIO object name
    metadata = await get_session_metadata(doc_key)
    if not metadata or "object_name" not in metadata:
        logger.warning("Document metadata missing for OnlyOffice download", key=doc_key)
        raise HTTPException(status_code=404, detail="Document session not found")
    if not _verify_download_token(doc_key, metadata, download_token):
        raise HTTPException(status_code=403, detail="Invalid download token")
    object_name = metadata["object_name"]
    
    docx_bytes = await _document_store.get(object_name)
    if not docx_bytes:
        logger.warning("Document not found in MinIO", object_name=object_name, key=doc_key)
        raise HTTPException(status_code=404, detail="Document not found")
    
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={doc_key}.docx",
        },
    )


class CallbackPayload(BaseModel):
    key: str
    status: int
    url: str | None = None
    users: list[str] | None = None
    actions: list[dict] | None = None
    changesurl: str | None = None
    history: dict | None = None
    token: str | None = None


@router.api_route("/callback", methods=["GET", "POST"])
async def onlyoffice_callback(
    request: Request,
    payload: CallbackPayload = None,
    db: AsyncSession = Depends(get_db),
):
    """
    OnlyOffice callback endpoint.
    Handles both POST (data) and GET (ping).
    """
    if request.method == "GET":
        return JSONResponse(status_code=405, content={"error": 1})
    
    if not payload:
        return JSONResponse(status_code=400, content={"error": 1})
    if not _validate_callback_token(payload.token, payload.key):
        logger.warning("Rejected OnlyOffice callback with invalid token", key=payload.key)
        return JSONResponse(status_code=401, content={"error": 1})

    logger.info("OnlyOffice callback received", key=payload.key, status=payload.status)
    
    # Status 2 = ready to save, 6 = force save
    if payload.status in (2, 6) and payload.url:
        metadata = await get_session_metadata(payload.key)
        if not metadata:
            logger.warning("Unknown document key in callback (Redis lookup failed)", key=payload.key)
            return JSONResponse(status_code=404, content={"error": 1})
            
        is_library_block = metadata.get("type") == "library_block"
        is_section = metadata.get("type") == "section"
        
        try:
            download_url = payload.url
            # OnlyOffice might send a URL with 'localhost:8443' if accessed via the host.
            # The backend container must use 'onlyoffice' (port 80) to reach it.
            if "localhost:8443" in download_url:
                download_url = download_url.replace("localhost:8443", "onlyoffice")
            elif "localhost" in download_url:
                download_url = download_url.replace("localhost", "onlyoffice")
            elif "127.0.0.1:8443" in download_url:
                download_url = download_url.replace("127.0.0.1:8443", "onlyoffice")
            elif "127.0.0.1" in download_url:
                download_url = download_url.replace("127.0.0.1", "onlyoffice")
            
            logger.info("Downloading updated document from OnlyOffice", url=download_url)
            
            # Download the updated document from OnlyOffice
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(download_url)
                response.raise_for_status()
                updated_docx = response.content
            
            # Update in MinIO store using stable object name
            object_name = metadata.get("object_name", payload.key)
            await _document_store.save(object_name, updated_docx)
            
            # Extract text from the updated .docx
            extracted_text = _extract_text_from_docx(updated_docx)
            
            logger.info(
                "Text extracted from OnlyOffice document",
                key=payload.key,
                text_length=len(extracted_text),
            )
            
            if is_section:
                proposal_id = metadata["proposal_id"]
                section_id = metadata["section_id"]
                # Update section content in database
                result = await db.execute(
                    select(ProposalSection).where(
                        ProposalSection.id == section_id,
                        ProposalSection.proposal_id == proposal_id,
                    )
                )
                section = result.scalar_one_or_none()
                if section:
                    section.content = _text_to_tiptap_content(extracted_text)
                    await db.flush()

                    proposal_result = await db.execute(select(Proposal).where(Proposal.id == proposal_id))
                    proposal = proposal_result.scalar_one_or_none()
                    if proposal:
                        compliance_events = await sync_requirement_compliance_and_gate(
                            db,
                            tender_id=proposal.tender_id,
                            actor_id=None,
                        )
                        await sync_tender_and_publish_event(
                            db,
                            tender_id=proposal.tender_id,
                            actor_id=None,
                            source="onlyoffice_callback",
                            event_type="proposal_section_updated",
                            event_payload=build_proposal_section_updated_event_payload(
                                section=section,
                                change_type="section_content_saved",
                                changed_fields=["content"],
                                source="onlyoffice",
                            ),
                        )
                        for event_type, event_payload in compliance_events:
                            await publish_domain_event(
                                db,
                                tender_id=proposal.tender_id,
                                actor_id=None,
                                source="onlyoffice_callback",
                                event_type=event_type,
                                payload=event_payload,
                            )

                    logger.info("Section content updated from OnlyOffice", section_id=section_id)
                
                # Index into RAG
                if extracted_text.strip():
                    try:
                        rag_engine = request.app.state.rag_engine
                        await rag_engine.ensure_initialized()
                        from app.ingestion.pipeline import IngestionPipeline
                        pipeline = IngestionPipeline(rag_engine)
                        
                        stats = await pipeline.ingest_text(
                            text=extracted_text,
                            document_id=section_id,
                            doc_type="proposal_section",
                            metadata={
                                "proposal_id": proposal_id,
                                "section_id": section_id,
                                "source": "onlyoffice",
                            },
                        )
                        
                        logger.info("Section content indexed in RAG", section_id=section_id, stats=stats)
                    except Exception as e:
                        logger.error("Failed to index section in RAG", error=str(e), section_id=section_id)
            
            elif is_library_block:
                block_id = metadata["block_id"]
                # Update library block content in database
                result = await db.execute(select(ContentBlock).where(ContentBlock.id == block_id))
                block = result.scalar_one_or_none()
                if block:
                    block.content = extracted_text
                    await db.flush()
                    
                    logger.info("Library block content updated from OnlyOffice", block_id=block_id)
                
                # Index into RAG
                if extracted_text.strip():
                    try:
                        rag_engine = request.app.state.rag_engine
                        await rag_engine.ensure_initialized()
                        from app.ingestion.pipeline import IngestionPipeline
                        pipeline = IngestionPipeline(rag_engine)
                        
                        stats = await pipeline.ingest_text(
                            text=extracted_text,
                            document_id=block_id,
                            doc_type="content_block",
                            metadata={
                                "block_id": block_id,
                                "category": block.category if block else None,
                                "source": "onlyoffice_library",
                            },
                        )
                        
                        logger.info("Library block indexed in RAG", block_id=block_id, stats=stats)
                    except Exception as e:
                        logger.error("Failed to index library block in RAG", error=str(e), block_id=block_id)

        except Exception as e:
            logger.error(
                "Failed to process OnlyOffice callback",
                error=str(e),
                key=payload.key,
            )
            return JSONResponse(status_code=500, content={"error": 1})
    
    # OnlyOffice expects {"error": 0} on success
    return {"error": 0}


@router.post("/forcesave/library/{block_id}")
async def force_save_library(
    block_id: int,
    current_user: UserResponse = Depends(get_current_user),
):
    """
    Trigger a force save in OnlyOffice for a library block.
    """
    # Find the active key for this block in Redis
    active_key = await get_active_key("block", str(block_id))
    
    if not active_key:
        raise HTTPException(status_code=404, detail="No active editing session found")
    
    try:
        command_url = f"{settings.onlyoffice_internal_url}/coauthoring/CommandService.ashx"
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.post(
                command_url,
                json={
                    "c": "forcesave",
                    "key": active_key,
                },
            )
            result = response.json()
            
        logger.info("Library force save triggered", key=active_key, result=result)
        return {"status": "ok", "result": result}
    except Exception as e:
        logger.error("Library force save failed", error=str(e))
        raise HTTPException(status_code=500, detail=f"Force save failed: {str(e)}")


@router.post("/forcesave/proposal/{proposal_id}/{section_id}")
async def force_save(
    proposal_id: int,
    section_id: int,
    current_user: UserResponse = Depends(get_current_user),
):
    """
    Trigger a force save in OnlyOffice.
    This calls the OnlyOffice Document Server command service to force saving.
    """
    # Find the active key for this section in Redis
    active_key = await get_active_key("section", f"{proposal_id}:{section_id}")
    
    if not active_key:
        raise HTTPException(status_code=404, detail="No active editing session found")
    
    try:
        command_url = f"{settings.onlyoffice_internal_url}/coauthoring/CommandService.ashx"
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.post(
                command_url,
                json={
                    "c": "forcesave",
                    "key": active_key,
                },
            )
            result = response.json()
            
        logger.info("Force save triggered", key=active_key, result=result)
        return {"status": "ok", "result": result}
    except Exception as e:
        logger.error("Force save failed", error=str(e))
        raise HTTPException(status_code=500, detail=f"Force save failed: {str(e)}")
